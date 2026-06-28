from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import joblib
import re
from io import BytesIO
from flask_mysqldb import MySQL
from werkzeug.utils import secure_filename
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
import numpy as np
import xgboost as xgb
import lime.lime_text
from sentence_transformers import SentenceTransformer, util
import wikipediaapi
import os
from dotenv import load_dotenv
from newsapi import NewsApiClient
import spacy
import requests
from reportlab.lib.pagesizes import A4
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Image
)
from reportlab.lib.styles import getSampleStyleSheet
import io
from reportlab.graphics import renderPDF
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor
from reportlab.lib.utils import ImageReader
from reportlab.graphics.charts.piecharts import Pie
from reportlab.graphics.shapes import (
    Drawing,
    Circle,
    String
)
import bcrypt
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_mail import Mail, Message
import random
from datetime import datetime, timedelta
import logging
from flask_wtf.csrf import CSRFProtect
from bs4 import BeautifulSoup
from newspaper import Article
from flask import send_from_directory
import fitz
import docx
from werkzeug.security import generate_password_hash
import mysql.connector
from google.cloud import storage


otp_storage = {}
verified_emails = set()
failed_login_attempts = {}
locked_accounts = {}
reset_otp_storage = {}

load_dotenv()

NEWS_API_KEY = os.getenv("NEWS_API_KEY")
GOOGLE_FACT_KEY = os.getenv("GOOGLE_FACT_KEY")
TRUSTED_SOURCES = [

    "Reuters",
    "BBC News",
    "Associated Press",
    "AP News",
    "The Hindu",
    "Indian Express",
    "NPR",
    "Al Jazeera English",
    "The Times of India",
    "Hindustan Times"

]

newsapi = NewsApiClient(api_key=NEWS_API_KEY)

app = Flask(__name__)
CORS(
    app,
    resources={
        r"/*": {
            "origins": "*"
        }
    }
)

storage_client = None
bucket = None

if os.getenv("K_SERVICE"):   # Running on Cloud Run
    storage_client = storage.Client()
    bucket = storage_client.bucket(
        os.getenv("PROFILE_BUCKET")
    )

# PROFILE_UPLOAD_FOLDER = 'profile_pictures'
# os.makedirs(PROFILE_UPLOAD_FOLDER, exist_ok=True)
# app.config['PROFILE_UPLOAD_FOLDER'] = PROFILE_UPLOAD_FOLDER

@app.after_request
def security_headers(response):

    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['X-XSS-Protection'] = '1; mode=block'

    return response

@app.errorhandler(413)
def too_large(e):

    return jsonify({
        "status": "error",
        "message": "File too large. Maximum allowed size is 10 MB."
    }), 413

# @app.errorhandler(Exception)
# def handle_exception(e):
#     logging.error(f"SERVER ERROR: {str(e)}")
#     return jsonify({
#         "status": "error",
#         "message": "Internal server error"
#     }), 500

@app.errorhandler(Exception)
def handle_exception(e):

    print("SERVER ERROR:")
    print(str(e))

    return jsonify({
        "status":"error",
        "message":str(e)
    }),500

logging.basicConfig(
    filename='security.log',
    level=logging.INFO,
    format='%(asctime)s %(levelname)s %(message)s'
)

def add_log(action, description):

    cursor = mysql.connection.cursor()

    cursor.execute(
        """
        INSERT INTO activity_logs
        (
            action,
            description
        )
        VALUES
        (
            %s,
            %s
        )
        """,
        (
            action,
            description
        )
    )

    mysql.connection.commit()

    cursor.close()

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024

app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = os.getenv("MAIL_USERNAME")
app.config['MAIL_PASSWORD'] = os.getenv("MAIL_PASSWORD")
app.config['MYSQL_CHARSET'] = 'utf8mb4'
app.config['SECRET_KEY'] = os.getenv("SECRET_KEY")
app.config['MAIL_DEFAULT_SENDER'] = 'support.verilexai@gmail.com'
csrf = CSRFProtect(app)
mail = Mail(app)

limiter = Limiter(
    key_func=get_remote_address,
    app=app,
    storage_uri="memory://"
)
# ============================================
# 🗄️ DATABASE CONFIG
# ============================================

app.config['MYSQL_HOST'] = os.getenv("MYSQL_HOST")
app.config['MYSQL_USER'] = os.getenv("MYSQL_USER")
app.config['MYSQL_PASSWORD'] = os.getenv("MYSQL_PASSWORD")
app.config['MYSQL_DB'] = os.getenv("MYSQL_DB")
app.config['MYSQL_PORT'] = int(os.getenv("MYSQL_PORT", 3306))
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10 MB max upload

mysql = MySQL(app)

# semantic_model = SentenceTransformer(
#     'all-MiniLM-L6-v2'
# )
semantic_model = None
def get_semantic_model():
    global semantic_model
    if semantic_model is None:
        semantic_model = SentenceTransformer(
            'all-MiniLM-L6-v2'
        )
    return semantic_model

nlp = spacy.load("en_core_web_sm")

wiki = wikipediaapi.Wikipedia(
    language='en',
    user_agent='VeriLexAI/1.0 (student-project)',
    extract_format=wikipediaapi.ExtractFormat.WIKI
)

# ============================================
# 📦 LOAD MODELS
# ============================================

lr_model = joblib.load('models/fake_news/lr_model.pkl')
tfidf_fake = joblib.load('models/fake_news/tfidf_vectorizer.pkl')

legal_model = joblib.load('models/legal/xgb_model.pkl')
tfidf_legal = joblib.load('models/legal/tfidf_vectorizer.pkl')

print("Loading fact verification model...")

# ============================================
# 🔍 XAI EXPLAINERS
# ============================================

lime_explainer = lime.lime_text.LimeTextExplainer(
    class_names=['Real News', 'Fake News']
)

# ============================================
# 📁 FILE SETTINGS
# ============================================

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}
ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}

ALLOWED_MIME_TYPES = {
    'text/plain',
    'application/pdf',
    'application/msword',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
}

def is_valid_mime_type(file):

    mime_type = file.content_type

    return mime_type in ALLOWED_MIME_TYPES

# def load_constitution_text():

#     try:

#         reader = PdfReader("constitution.pdf")

#         full_text = ""

#         for page in reader.pages:

#             text = page.extract_text()

#             if text:
#                 full_text += text + "\n"

#         return full_text

#     except Exception as e:

#         print(
#             "Constitution PDF Error:",
#             str(e)
#         )

#         return ""
# constitution_text = load_constitution_text()

# ============================================
# 🧹 TEXT PREPROCESSING
# ============================================

def clean_text(text):
    text = text.lower()
    text = re.sub(r'[^a-zA-Z ]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def extract_claims(text):
    doc = nlp(text)
    claims = []
    for sent in doc.sents:
        sentence = sent.text.strip()
        if (len(sentence) > 25 and any(token.pos_ == "VERB" for token in nlp(sentence))):
            claims.append(sentence)
    return claims[:15]

def extract_key_entities(text):
    doc = nlp(text)
    entities = []
    for ent in doc.ents:
        if ent.label_ in ["PERSON", "ORG", "GPE"]:
            entities.append(
                ent.text
            )
    return list(set(entities))

def is_valid_email(email):
    pattern = r'^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$'
    return re.match(pattern, email)

def is_valid_password(password):

    pattern = (
        r'^(?=(?:.*[A-Z]){2,})'
        r'(?=(?:.*[a-z]){2,})'
        r'(?=(?:.*\d){2,})'
        r'(?=.*[!@#$%^&*()_+\-=\[\]{};\'":\\|,.<>\/?])'
        r'.{10,}$'
    )

    return re.match(pattern, password)

def verify_with_google_fact_check(claim):

    try:

        url = (
            "https://factchecktools.googleapis.com/v1alpha1/"
            "claims:search"
        )

        response = requests.get(
            url,
            params={
                "query": claim,
                "key": GOOGLE_FACT_KEY
            },
            timeout=10
        )


        data = response.json()


        claims = data.get(
            "claims",
            []
        )


        if not claims:

            return None


        reviews = claims[0].get(
            "claimReview",
            []
        )


        if not reviews:

            return None


        rating = reviews[0].get(
            "textualRating",
            ""
        ).lower()


        print("Google rating:", rating)


        if "true" in rating or "correct" in rating:

            return (
                True,
                f"Google Fact Check: {rating}"
            )


        if (
            "false" in rating
            or "fake" in rating
            or "incorrect" in rating
        ):

            return (
                False,
                f"Google Fact Check: {rating}"
            )


        return None


    except Exception as e:

        print(
            "Google API Error:",
            e
        )

        return None
    
def predict_ml_fake_news(text):

    cleaned = clean_text(text)

    vector = tfidf_fake.transform([cleaned])

    prediction = lr_model.predict(vector)[0]

    probability = lr_model.predict_proba(vector)[0]

    confidence = round(max(probability) * 100)

    return prediction, confidence
    

def verify_with_gemini(claim):

    try:

        key = os.getenv("GEMINI_API_KEY")

        if not key:

            print("Gemini API key missing")

            return None

        genai.configure(
            api_key=key
        )

        model = genai.GenerativeModel(
            "gemini-2.5-flash"
        )

        prompt = f"""
You are an expert fact-checking assistant.

Analyze ONLY the factual accuracy of the following claim.

Claim:
{claim}

Return ONLY:

VERDICT: TRUE

or

VERDICT: FALSE

or

VERDICT: UNCERTAIN

REASON:
One short sentence.
"""

        response = model.generate_content(
            prompt
        )

        answer = response.text.strip()

        upper = answer.upper()

        if "VERDICT: TRUE" in upper:

            return True, answer

        elif "VERDICT: FALSE" in upper:

            return False, answer

        return None, answer

    except Exception as e:

        print(
            "Gemini verification error:",
            str(e)
        )

        return (
        None,
        "Gemini unavailable"
    )

def extract_claim_parts(text):

    doc = nlp(text)

    persons = []
    places = []
    nouns = []

    for ent in doc.ents:

        if ent.label_ == "PERSON":
            persons.append(ent.text)

        elif ent.label_ in ["GPE", "LOC"]:
            places.append(ent.text)

    for token in doc:

        if token.pos_ in ["NOUN", "PROPN"]:

            nouns.append(
                token.text.lower()
            )

    return persons, places, nouns


def search_live_news(query):

    # ---------- NewsAPI ----------
    try:

        articles = newsapi.get_everything(
            q=query[:100],
            language='en',
            sort_by='relevancy',
            page_size=5
        )

        results = articles.get(
            "articles",
            []
        )

        if results:

            print(
                "NewsAPI articles:",
                len(results)
            )

            return results

    except Exception as e:

        print(
            "News API Error:",
            e
        )

    # ---------- NewsData.io ----------
    try:

        url = (
            "https://newsdata.io/api/1/news"
        )

        params = {

            "apikey":
            os.getenv(
                "NEWSDATA_API_KEY"
            ),

            "q":
            query,

            "language":
            "en"

        }

        r = requests.get(
            url,
            params=params,
            timeout=10
        )

        data = r.json()

        results = []

        for item in data.get(
            "results",
            []
        )[:5]:

            results.append({

                "title":
                item.get(
                    "title",
                    ""
                ),

                "description":
                item.get(
                    "description",
                    ""
                ),

                "source":
                {
                    "name":
                    item.get(
                        "source_id",
                        ""
                    )
                },

                "url":
                item.get(
                    "link",
                    ""
                )

            })

        print(
            "NewsData articles:",
            len(results)
        )

        return results

    except Exception as e:

        print(
            "NewsData Error:",
            e
        )

        return []

def crawl_news_articles(query):

    try:

        articles = search_live_news(query)

        extracted_articles = []

        for item in articles:

            url = item.get("url")

            if not url:
                continue

            try:

                article = Article(url)

                article.download()

                article.parse()

                content = article.text[:3000]

                extracted_articles.append({

                    "title": item.get("title"),

                    "source": item.get("source", {}).get("name"),

                    "url": url,

                    "content": content

                })

            except Exception as e:

                print("Article extraction failed:", e)

        return extracted_articles

    except Exception as e:

        print("Crawler Error:", e)

        return []
    
def check_trusted_news_sources(text):

    try:

        articles = search_live_news(text)

        if not articles:
            return 0, []

        score = 0
        sources_found = []

        for article in articles:

            source_name = article.get(
                "source",
                {}
            ).get(
                "name",
                ""
            )

            if source_name in TRUSTED_SOURCES:

                score += 20
                sources_found.append(source_name)

        score = min(score, 100)

        return score, list(set(sources_found))

    except Exception as e:

        print("Source verification error:", e)

        return 0, []
    
def generate_fast_evidence(text):

    try:

        entities = extract_key_entities(text)

        if entities:

            query = " ".join(entities[:3])

        else:

            query = text[:80]

        articles = search_live_news(query)
        print("QUERY =", query)
        print("ARTICLES FOUND =", len(articles))

        if not articles:

            return {
                "summary": "No recent evidence found.",
                "sources": []
            }

        titles = []
        sources = []

        for article in articles[:5]:

            title = article.get("title", "")

            source = article.get(
                "source",
                {}
            ).get(
                "name",
                ""
            )

            if title:
                titles.append(title)

            if source:
                sources.append(source)
        print("SUMMARY =", " | ".join(titles))
        print("SOURCES =", list(set(sources)))

        return {

            "summary":
            " | ".join(titles),

            "sources":
            list(set(sources))
        }

    except Exception as e:

        print("Evidence Error:", e)

        return {
            "summary": "",
            "sources": []
        }
    
def verify_with_live_news(text):
    model = get_semantic_model()

    try:

        articles = crawl_news_articles(text)

        if not articles:

            return 0

        claim_embedding = model.encode(
            text,
            convert_to_tensor=True
        )

        similarities = []

        for article in articles:

            headline = article.get(
                "title",
                ""
            )

            description = article.get(
                "description",
                ""
            )

            content = (
                headline +
                " " +
                description
            )

            if not content:
                continue

            article_embedding = model.encode(
                content,
                convert_to_tensor=True
            )

            similarity = util.cos_sim(
                claim_embedding,
                article_embedding
            ).item()

            similarities.append(similarity)

        if not similarities:
            return 0

        return round(
            max(similarities) * 100
        )

    except Exception as e:

        print(
            "Semantic verification error:",
            e
        )

        return 0
    
def generate_evidence_analysis(text):

    try:

        doc = nlp(text)

        subjects = []

        for token in doc:

            if token.pos_ in ["NOUN", "PROPN"]:
                subjects.append(token.text)

        if not subjects:

            return {
                "summary": "No reliable evidence found.",
                "sources": []
            }

        topic = text[:100]

        page = wiki.page(topic)

        if page.exists():

            summary = page.summary[:500]

            return {
                "summary": summary,
                "sources": ["Wikipedia"]
            }

        # FALLBACK TO NEWS API
        articles = newsapi.get_everything(
            q=topic,
            language='en',
            sort_by='relevancy',
            page_size=3
        )

        article_list = articles.get("articles", [])

        if not article_list:

            return {
                "summary": "No reliable evidence found.",
                "sources": []
            }

        evidence_text = ""

        sources = []

        for art in article_list:

            title = art.get("title", "")

            source = art.get("source", {}).get("name", "")

            evidence_text += f"{title}. "

            sources.append(source)

        return {
            "summary": evidence_text,
            "sources": list(set(sources))
        }

    except Exception as e:

        print("Evidence Error:", e)

        return {
            "summary": "Evidence unavailable.",
            "sources": []
        }

# ============================================
# 🤖 MODEL FUNCTIONS
# ============================================

# def check_local_fact_database(claim):

#     try:

#         cur = mysql.connection.cursor()

#         # cur.execute(
#         #     """
#         #     SELECT claim_text, verdict
#         #     FROM fact_database
#         #     """
#         # )

#         rows = cur.fetchall()

#         cur.close()

#         if not rows:
#             return None

#         claim_embedding = semantic_model.encode(
#             claim,
#             convert_to_tensor=True
#         )

#         best_similarity = 0
#         best_verdict = None

#         for row in rows:

#             db_claim = row[0]
#             verdict = row[1]

#             db_embedding = semantic_model.encode(
#                 db_claim,
#                 convert_to_tensor=True
#             )

#             similarity = util.cos_sim(
#                 claim_embedding,
#                 db_embedding
#             ).item()

#             if similarity > best_similarity:

#                 best_similarity = similarity
#                 best_verdict = verdict

#         if best_similarity >= 0.75:

#             return best_verdict

#         return None

#     except Exception as e:

#         print(
#             "Fact DB Error:",
#             e
#         )

#         return None

def check_local_fact_database(claim):

    try:

        cur = mysql.connection.cursor()

        cur.execute(
            """
            SELECT verdict
            FROM fact_database
            WHERE claim_text LIKE %s
            LIMIT 1
            """,
            (f"%{claim[:100]}%",)
        )

        row = cur.fetchone()

        cur.close()

        if row:

            return row[0]

        return None

    except Exception as e:

        print(
            "Fact DB Error:",
            e
        )

        return None

def basic_fact_checks(text):

    text = text.lower()

    known_false = [

        "moon is the satellite of mars",

        "earth is flat",

        "sun revolves around earth",

        "india is in europe"

    ]

    known_true = [

        "modi is the prime minister of india",

        "earth revolves around sun",

        "water boils at 100 degrees celsius"

    ]

    for fact in known_false:

        if fact in text:

            return (
                "🟥 Fake News",
                "Known factual contradiction",
                99,
                1,
                99
            )

    for fact in known_true:

        if fact in text:

            return (
                "🟩 Real News",
                "Known verified fact",
                99,
                99,
                1
            )

    return None


def predict_fake_news(text):

    # -------------------------
    # EXTRACT CLAIMS
    # -------------------------

    claims = extract_claims(text)

    if not claims:
        claims = [text]

    claims = claims[:3]

    rule_result = basic_fact_checks(text)

    entities = extract_key_entities(text)

    if rule_result is not None:
        return rule_result

    print("Checking:", text)

    real_score = 0

    fake_score = 0
    explanation_parts = []

    try:

        ml_prediction, ml_confidence = \
            predict_ml_fake_news(text)

        if ml_prediction == 1:

            real_score += 50

            explanation_parts.append(
                f"ML model predicts real "
                f"({ml_confidence}%)"
            )

        else:

            fake_score += 50

            explanation_parts.append(
                f"ML model predicts fake "
                f"({ml_confidence}%)"
            )

    except Exception as e:

        print("ML Error:", e)

    skip_gemini = False

    # -------------------------
    # ENTITY VERIFICATION
    # -------------------------

    # for entity in entities:

    #     articles = search_live_news(entity)

    #     if articles:

    #         real_score += 2

    #         explanation_parts.append(

    #             f"Entity found in news: {entity}"

    #         )

    # -------------------------
    # LOCAL DB + GOOGLE FACT CHECK
    # -------------------------

    for claim in claims:

        local_result = check_local_fact_database(

            claim

        )

        if local_result:

            skip_gemini = True

            if "true" in local_result.lower():

                real_score += 80

                explanation_parts.append(

                    "Matched local fact database"

                )

            elif "false" in local_result.lower():

                fake_score += 80

                explanation_parts.append(

                    "Matched local fact database"

                )

            continue

        google_result = (

            verify_with_google_fact_check(

                claim

            )

        )

        if google_result:

            is_true, explanation = google_result

            explanation_parts.append(

                explanation

            )

            if is_true:

                real_score += 60

            else:

                fake_score += 60

    # -------------------------
    # GEMINI FACT VERIFICATION
    # -------------------------

    if claims and not skip_gemini:

        merged_claim = ". ".join(

            claims

        )

        gemini_result = (

            verify_with_gemini(

                merged_claim

            )

        )

        if gemini_result:

            gemini_verdict, gemini_reason = (

                gemini_result

            )

            explanation_parts.append(

                "Gemini: "

                + gemini_reason

            )

            if gemini_verdict is True:

                real_score += 35

            elif gemini_verdict is False:

                fake_score += 35

            else:

                real_score += 5

                fake_score += 5
    for claim in claims:

        gemini_result = verify_with_gemini(claim)

        if not gemini_result:
            continue

        verdict, reason = gemini_result

        explanation_parts.append(
        f"Gemini: {reason}"
    )

        if verdict is True:
            real_score += 15

        elif verdict is False:
            fake_score += 15

    # -------------------------
    # TRUSTED NEWS SOURCES
    # -------------------------

    source_score, sources = (

        check_trusted_news_sources(

            text

        )

    )

    if source_score > 0:

        real_score += source_score

        explanation_parts.append(

            f"Found in trusted sources: "

            f"{', '.join(sources)}"

        )

    # -------------------------
    # LIVE NEWS MATCHING
    # -------------------------

    semantic_score = (verify_with_live_news(text))

    if semantic_score >= 75:

        real_score += 30

        explanation_parts.append(

            f"Strong match with live news "

            f"coverage "

            f"({semantic_score}%)"

        )

    elif semantic_score <= 35:

        explanation_parts.append("No strong supporting evidence found.")

        explanation_parts.append(

            "Limited recent news "

            "evidence available."

        )

        explanation_parts.append(

            f"Little supporting "

            f"evidence found "

            f"({semantic_score}%)"

        )

    # -------------------------
    # FINAL DECISION
    # -------------------------

    total = (

        real_score

        + fake_score

    )

    if total == 0:

        return (

            "🟨 Unverified",

            "Insufficient evidence found",

            50,

            50,

            50

        )

    if abs(

        real_score

        - fake_score

    ) <= 20:

        return (

            "🟨 Unverified",

            "Conflicting evidence found",

            55,

            real_score,

            fake_score

        )

    confidence = round(

        max(

            real_score,

            fake_score

        )

        / total

        * 100

    )

    if real_score >= fake_score:

        return (

            "🟩 Likely Real",

            " | ".join(

                explanation_parts

            ),

            confidence,

            real_score,

            fake_score

        )

    return (

        "🟥 Likely Fake",

        " | ".join(

            explanation_parts

        ),

        confidence,

        real_score,

        fake_score

    )



def predict_legal(text):

    cleaned = clean_text(text)

    vector = tfidf_legal.transform(
        [cleaned]
    )


    prediction = legal_model.predict(
        vector
    )[0]


    probabilities = legal_model.predict_proba(
        vector
    )[0]


    confidence = round(
        max(probabilities) * 100
    )


    result_text = (

        "🟢 Bail Granted"

        if prediction == 1

        else "🔴 Bail Denied"

    )


    return (

        int(prediction),

        result_text,

        confidence

    )


def save_history(email, analysis_type, input_text, result):

    try:
        print("INSIDE SAVE_HISTORY")
        cur = mysql.connection.cursor()

        cur.execute(
            """
            INSERT INTO analysis_history
            (user_email, analysis_type, input_text, result)

            VALUES (%s,%s,%s,%s)
            """,
            (
                email,
                analysis_type,
                input_text,
                result
            )
        )

        mysql.connection.commit()
        print("INSERT SUCCESS")
        cur.close()

    except Exception as e:

        print(
            "History save error:",
            str(e)
        )

# ============================================
# 🧠 XAI HELPERS
# ============================================

def explain_legal_shap(text):
    cleaned = clean_text(text)
    vector = tfidf_legal.transform([cleaned])

    # Use XGBoost built-in SHAP-style contributions
    booster = legal_model.get_booster() if hasattr(legal_model, "get_booster") else legal_model
    dmatrix = xgb.DMatrix(vector)

    shap_contribs = booster.predict(dmatrix, pred_contribs=True)

    # Last value is bias term, so remove it
    contrib_values = shap_contribs[0][:-1]

    feature_names = tfidf_legal.get_feature_names_out()
    dense_vector = vector.toarray()[0]
    active_indices = np.where(dense_vector > 0)[0]

    feature_importance = []
    for idx in active_indices:
        feature_importance.append((feature_names[idx], contrib_values[idx]))

    feature_importance = sorted(
        feature_importance,
        key=lambda x: abs(x[1]),
        reverse=True
    )

    if feature_importance:
        top_terms = ', '.join([f'"{feat}"' for feat, _ in feature_importance[:6]])
        return f'Important legal terms influencing the prediction: {top_terms}.'
    return 'The model used multiple legal-text patterns to make this prediction.'

# ============================================
# 📂 FILE HELPERS
# ============================================

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_image(filename):
    return (
        '.' in filename and
        filename.rsplit('.', 1)[1].lower()
        in ALLOWED_IMAGE_EXTENSIONS
    )

def extract_text_from_txt(file_storage):
    file_storage.seek(0)
    return file_storage.read().decode('utf-8', errors='ignore').strip()

def extract_text_from_pdf(file_storage):
    file_storage.seek(0)
    pdf_stream = BytesIO(file_storage.read())
    reader = PdfReader(pdf_stream)

    extracted_pages = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text and page_text.strip():
            extracted_pages.append(page_text.strip())

    return '\n'.join(extracted_pages).strip()

def extract_text_from_docx(file_storage):
    file_storage.seek(0)
    doc_stream = BytesIO(file_storage.read())
    document = Document(doc_stream)
    paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
    return '\n'.join(paragraphs).strip()

def extract_text_from_doc(file_storage):
    raise ValueError("Old .doc files are not supported yet. Please convert the file to .docx, .pdf, or .txt.")

def extract_text_from_file(file_storage):

    filename = secure_filename(file_storage.filename)

    extension = filename.rsplit('.', 1)[1].lower()

    if extension == 'txt':
        return extract_text_from_txt(file_storage)

    elif extension == 'pdf':
        return extract_text_from_pdf(file_storage)

    elif extension == 'docx':
        return extract_text_from_docx(file_storage)

    elif extension == 'doc':
        return extract_text_from_doc(file_storage)

    else:
        raise ValueError("Unsupported file type.")
    


def generate_analysis_report(
        input_text,
        prediction,
        confidence,
        explanation):

    from reportlab.pdfbase.pdfmetrics import stringWidth

    input_text = str(input_text or "No input")
    prediction = str(prediction or "No prediction")
    explanation = str(explanation or "No explanation")

    try:
        confidence = int(float(confidence))
    except:
        confidence = 0

    pdf_buffer = io.BytesIO()

    PAGE_WIDTH = 1536
    PAGE_HEIGHT = 1002

    c = canvas.Canvas(
        pdf_buffer,
        pagesize=(PAGE_WIDTH, PAGE_HEIGHT)
    )

    # =========================================
    # BACKGROUND TEMPLATE
    # =========================================

    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

    template_path = os.path.join(
    BASE_DIR,
    "assets",
    "report_template.png"
)

    if not os.path.exists(template_path):
        raise FileNotFoundError(
            f"Template not found: {template_path}"
        )

    bg = ImageReader(template_path)

    c.drawImage(
        bg,
        0,
        0,
        width=PAGE_WIDTH,
        height=PAGE_HEIGHT
    )

    # =========================================
    # COMMON POSITIONS
    # =========================================

    VALUE_X = 430

    # =========================================
    # DATE
    # =========================================

    c.setFillColor(HexColor("#111827"))

    c.setFont(
        "Helvetica",
        24
    )

    formatted_date = datetime.now().strftime(
        "%d %B %Y at %I:%M %p"
    )

    c.drawString(
        VALUE_X,
        700,
        formatted_date
    )

    # =========================================
    # PREDICTION
    # =========================================

    clean_prediction = prediction.replace("🟢", "") \
                                 .replace("🔴", "") \
                                 .replace("🟥", "") \
                                 .replace("🟩", "") \
                                 .strip()

    if (
        "Fake" in prediction
        or "Denied" in prediction
    ):
        status_color = HexColor("#EF4444")

    else:
        status_color = HexColor("#10B981")

    # STATUS DOT
    c.setFillColor(status_color)

    c.circle(
        VALUE_X + 12,
        585,
        10,
        fill=1
    )

    # PREDICTION TEXT
    c.setFillColor(HexColor("#111827"))

    c.setFont(
        "Helvetica-Bold",
        30
    )

    c.drawString(
        VALUE_X + 40,
        575,
        clean_prediction
    )

    # =========================================
    # CONFIDENCE
    # =========================================

    if confidence >= 70:
        confidence_color = HexColor("#10B981")

    elif confidence >= 40:
        confidence_color = HexColor("#3B82F6")

    else:
        confidence_color = HexColor("#EF4444")

    c.setFillColor(confidence_color)

    c.setFont(
        "Helvetica-Bold",
        38
    )

    c.drawString(
        VALUE_X,
        450,
        f"{confidence}%"
    )

    # =========================================
    # INPUT TEXT
    # =========================================

    c.setFillColor(HexColor("#111827"))

    c.setFont(
        "Helvetica",
        23
    )

    # Proper alignment
    input_y = 340

    trimmed_input = input_text[:55]

    c.drawString(
        VALUE_X,
        input_y,
        trimmed_input
    )

    # =========================================
    # EXPLANATION
    # =========================================

    c.setFillColor(HexColor("#111827"))

    explanation_font_size = 18

    c.setFont(
        "Helvetica",
        explanation_font_size
    )

    max_width = 760

    words = explanation.split()

    lines = []
    current_line = ""

    for word in words:

        test_line = current_line + word + " "

        width = stringWidth(
            test_line,
            "Helvetica",
            explanation_font_size
        )

        if width < max_width:
            current_line = test_line

        else:
            lines.append(current_line)
            current_line = word + " "

    lines.append(current_line)

    # Better vertical alignment
    y_position = 220

    for line in lines[:3]:

        c.drawString(
            VALUE_X,
            y_position,
            line.strip()
        )

        y_position -= 26

    # =========================================
    # DONUT CHART
    # =========================================

    # MOVED UPWARD
    chart_x = 860
    chart_y = 345

    chart_width = 300
    chart_height = 300

    drawing = Drawing(
        chart_width,
        chart_height
    )

    pie = Pie()

    pie.x = 25
    pie.y = 25

    pie.width = 250
    pie.height = 250

    confidence = max(0, min(100, int(confidence)))
    remaining = 100 - confidence

    pie.data = [
        confidence,
        remaining
    ]

    pie.labels = ["", ""]

    pie.slices.strokeWidth = 1

    # MAIN COLOR
    if confidence >= 70:
        pie.slices[0].fillColor = HexColor("#10B981")

    elif confidence >= 40:
        pie.slices[0].fillColor = HexColor("#3B82F6")

    else:
        pie.slices[0].fillColor = HexColor("#EF4444")

    # REMAINING COLOR
    pie.slices[1].fillColor = HexColor("#E5E7EB")

    drawing.add(pie)

    # DONUT HOLE
    hole = Circle(
        150,
        150,
        70
    )

    hole.fillColor = HexColor("#FFFFFF")
    hole.strokeColor = HexColor("#FFFFFF")

    drawing.add(hole)

    # CENTER TEXT
    center_text = String(
        150,
        145,
        f"{confidence}%",
        textAnchor="middle"
    )

    center_text.fontName = "Helvetica-Bold"
    center_text.fontSize = 30
    center_text.fillColor = HexColor("#111827")

    drawing.add(center_text)

    renderPDF.draw(
        drawing,
        c,
        chart_x,
        chart_y
    )

    # =========================================
    # SAVE PDF
    # =========================================

    c.save()

    pdf_buffer.seek(0)

    return pdf_buffer

@csrf.exempt
@app.route("/generate-report", methods=["POST"])
def generate_report():

    try:

        data = request.get_json(force=True)

        pdf_buffer = generate_analysis_report(

            input_text=data.get("text", ""),

            prediction=data.get("result", ""),

            confidence=data.get("confidence", 0),

            explanation=data.get(
                "explanation",
                "No explanation available"
            )

        )

        pdf_buffer.seek(0)

        return send_file(

            pdf_buffer,

            mimetype="application/pdf",

            as_attachment=True,

            download_name="VeriLex_Report.pdf"

        )

    except Exception as e:

        return jsonify({

            "status": "error",

            "message": str(e)

        }), 500

# ============================================
# 🏠 HOME
# ============================================

@app.route('/', methods=['GET'])
def home():
    return jsonify({
        "message": "VeriLex AI Backend is Running 🚀"
    })

# ============================================
# 🔐 REGISTER
# ============================================

@csrf.exempt
@app.route('/register', methods=['POST'])
@limiter.limit("3 per minute")
def register():
    data = request.get_json()
    captcha_token = data.get("captcha")

    if not verify_recaptcha(captcha_token):
        return jsonify({
        "status": "error",
        "message": "CAPTCHA verification failed"
    }), 400
    name = data.get('name')
    email = data.get('email')
    password = data.get('password')
    if not is_valid_email(email):
        return jsonify({
        "status": "error",
        "message": "Invalid email format"
    }), 400
    if email not in verified_emails:
        return jsonify({
        "status": "error",
        "message": "Please verify OTP first."
    }), 403

    try:
        cur = mysql.connection.cursor()
        cur.execute("SELECT * FROM users WHERE email=%s", (email,))
        existing_user = cur.fetchone()

        if existing_user:
            cur.close()
            return jsonify({
                "status": "error",
                "message": "User already exists. Please login."
            }), 409

        hashed_password = bcrypt.hashpw(
        password.encode('utf-8'),
        bcrypt.gensalt()
        )
        cur.execute(
            "INSERT INTO users(name, email, password) VALUES(%s, %s, %s)",
            (name, email, hashed_password.decode('utf-8'))
        )
        mysql.connection.commit()
        cur.close()
        verified_emails.discard(email)
        return jsonify({
            "status": "success",
            "message": "Registration successful! Please login again."
        })

    except Exception as e:
        logging.error(str(e))
        return jsonify({
        "status": "error",
        "message": "Internal server error"
    }), 500

# ============================================
# 🔐 LOGIN
# ============================================

@csrf.exempt
@app.route('/login', methods=['POST'])
@limiter.limit("5 per minute")
def login():

    data = request.get_json()
    captcha_token = data.get("captcha")

    if not verify_recaptcha(captcha_token):
        return jsonify({
        "status": "error",
        "message": "CAPTCHA verification failed"
    }), 400
    email = data.get('email', '').strip()
    password = data.get('password', '').strip()

    try:

        # ACCOUNT LOCK CHECK
        if email in locked_accounts:

            lock_time = locked_accounts[email]

            if datetime.now() < lock_time:

                remaining = (
                    lock_time - datetime.now()
                ).seconds // 60

                return jsonify({
                    "status": "error",
                    "message":
                    f"Account locked. Try again in {remaining} minutes."
                }), 403

            else:

                del locked_accounts[email]
                failed_login_attempts[email] = 0

        cur = mysql.connection.cursor()

        cur.execute(
            "SELECT * FROM users WHERE email=%s",
            (email,)
        )

        user = cur.fetchone()

        cur.close()

        if not user:

            logging.warning(f"Login failed: User not found ({email})")

            return jsonify({
                "status": "error",
                "message": "User not found. Please register first."
            }), 404
        is_active = user[6]

        if user[6] == 0:

            logging.warning(
                f"Disabled account login attempt: {email}"
            )

            return jsonify({

                "status":"error",
                "message":
                "Your account has been disabled by the administrator."
            }), 403

        db_password = user[3]

        if not bcrypt.checkpw(
            password.encode('utf-8'),
            db_password.encode('utf-8')
        ):

            failed_login_attempts[email] = (
                failed_login_attempts.get(email, 0) + 1
            )

            logging.warning(
                f"Incorrect password for {email}"
            )

            if failed_login_attempts[email] >= 5:

                locked_accounts[email] = (
            datetime.now() + timedelta(minutes=20)
                )

                logging.warning(
                    f"Account locked: {email}"
                )

                return jsonify({
                    "status": "error",
                    "message":
                    "Too many failed attempts. Account locked for 20 minutes."
                }), 403

            remaining_attempts = 5 - failed_login_attempts[email]
            return jsonify({
                "status": "error",
                "message": f"Incorrect password. {remaining_attempts}/5 attempts remaining."
            }), 401

        # RESET FAILED ATTEMPTS
        failed_login_attempts[email] = 0

        logging.info(f"Successful login: {email}")

        cur = mysql.connection.cursor()

        cur.execute(
            """
            UPDATE users
            SET last_login = NOW()
            WHERE email=%s
            """,
            (email,)
        )

        mysql.connection.commit()

        cur.close()

        # CHECK IF PROFILE EXISTS

        cur = mysql.connection.cursor()

        cur.execute("""
        SELECT
        username,
        gender,
        dob,
        country,
        occupation,
        bio
        FROM user_profiles
        WHERE user_email=%s
        """, (email,))

        profile = cur.fetchone()

        if not profile:

            profile_completed = False

        else:

            username, gender, dob, country, occupation, bio = profile

            profile_completed = all([
                username,
                gender,
                dob,
                country,
                occupation,
                bio
            ])
        cur.close()
        
        return jsonify({
            "status": "success",
            "message": "Login successful",
            "name": user[1],
            "email": user[2],
            "profile_completed": profile_completed
        })

    except Exception as e:

        logging.error(f"LOGIN ERROR: {str(e)}")

        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500


def verify_recaptcha(token):

    secret_key = os.getenv("RECAPTCHA_SECRET")

    response = requests.post(
        "https://www.google.com/recaptcha/api/siteverify",
        data={
            "secret": secret_key,
            "response": token
        }
    )

    result = response.json()

    return result.get("success", False)

# ============================================
# 📰 FACT CHECK TEXT API
# ============================================

@csrf.exempt
@app.route('/fact-check', methods=['POST'])
def fact_check():

    try:
        data = request.get_json()

        text = data.get('text')
        # use_xai = data.get('use_xai', False)
        email = data.get('email')

        if not text or not text.strip():
            return jsonify({
                "status": "error",
                "message": "No text provided"
            }), 400
        
        prediction, live_explanation, confidence, real_confidence, fake_confidence = predict_fake_news(text)

        # default explanation
        explanation = live_explanation

        # if user wants XAI
        # if use_xai:
        #     try:
        #         explanation = explain_fake_news_lime(text)

        #     except Exception as e:
        #         explanation = (
        #             f"Explanation could not be generated: {str(e)}"
        #         )

        if email:
            print("EMAIL RECEIVED:", email)
            save_history(
                email,
                "Fact Check",
                text,
                prediction
            )
            print("SAVE HISTORY CALLED")

        # LIVE EVIDENCE
        evidence = generate_fast_evidence(text)
        print("========== EVIDENCE ==========")
        print(evidence)
        print("==============================")
        # LEGAL WARNING
        legal_warning = ""

        if "fake" in prediction.lower():
            legal_warning = (
                "This content may involve misinformation "
                "or defamation-related legal concerns."
            )

        return jsonify({
            "status": "success",
            "input_text": text,
            "result": prediction,
            "explanation": explanation,
            "confidence": confidence,
            "real_confidence": real_confidence,
            "fake_confidence": fake_confidence,
            "live_evidence": evidence["summary"],
            "sources": evidence["sources"],
            "legal_warning": legal_warning
        })

    except Exception as e:

        print("FACT CHECK ERROR:", str(e))

        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500


# ============================================
# 📰 FACT CHECK FILE API
# ============================================

@csrf.exempt
@app.route('/fact-check-upload', methods=['POST'])
def fact_check_upload():

    try:

        if 'file' not in request.files:
            return jsonify({
                "status": "error",
                "message": "No file uploaded."
            }), 400


        file = request.files['file']
        email = request.form.get('email')

        if file.filename == '':
            return jsonify({
                "status": "error",
                "message": "No file selected."
        }), 400

        if not is_valid_mime_type(file):
            return jsonify({
                "status": "error",
                "message": "Invalid file type detected."
        }), 400


        if not allowed_file(file.filename):
            return jsonify({
                "status": "error",
                "message":
                "Unsupported file type."
            }), 400


        extracted_text = extract_text_from_file(file)


        if not extracted_text.strip():
            return jsonify({
                "status": "error",
                "message":
                "Could not extract text."
            }), 400


        prediction, live_explanation, confidence, real_confidence, fake_confidence = predict_fake_news(extracted_text)

        explanation = live_explanation

        email = request.form.get("email")
        if email:
            save_history(
                email,
                "Fact Check",
                extracted_text,
                prediction
            )
        evidence = generate_fast_evidence(extracted_text)

        return jsonify({

            "status": "success",

            "filename": file.filename,

            "extracted_text": extracted_text[:3000],

            "result": prediction,

            "explanation": explanation,

            "confidence": confidence,

            "real_confidence": real_confidence,

            "fake_confidence": fake_confidence,

            "live_evidence": evidence["summary"],

            "sources": evidence["sources"],

            "legal_warning": ""
        })


    except Exception as e:

        print("UPLOAD ERROR:", str(e))

        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500
    
# ============================================
# ⚖️ LEGAL TEXT API
# ============================================

@csrf.exempt
@app.route('/legal', methods=['POST'])
def legal():
    data = request.get_json()
    text = data.get('text')
    use_xai = data.get('use_xai', False)
    email = data.get('email')

    if not text or not text.strip():
        return jsonify({
            "status": "error",
            "message": "No text provided"
        }), 400

    prediction, result_text, confidence = \
    predict_legal(text)
    explanation = ""

    if use_xai:
        try:
            explanation = explain_legal_shap(text)
        except Exception as e:
            explanation = f"Explanation could not be generated: {str(e)}"

    if email:
        print("EMAIL RECEIVED:", email)
        save_history(
        email,
        "Legal Analysis",
        text,
        result_text
    )
        print("SAVE HISTORY CALLED")

    return jsonify({
        "status": "success",
        "input_text": text,
        "prediction": prediction,
        "result": result_text,
        "confidence": confidence,
        "explanation": explanation
    })

# ============================================
# ⚖️ LEGAL FILE API
# ============================================

@csrf.exempt
@app.route('/legal-upload', methods=['POST'])
def legal_upload():
    if 'file' not in request.files:
        return jsonify({
            "status": "error",
            "message": "No file uploaded."
        }), 400

    file = request.files['file']
    email = request.form.get('email')
    use_xai = request.form.get('use_xai', 'false').lower() == 'true'

    if file.filename == '':
        return jsonify({
        "status": "error",
        "message": "No file selected."
    }), 400

    if not is_valid_mime_type(file):

        return jsonify({
        "status": "error",
        "message": "Invalid file type detected."
    }), 400

    if not allowed_file(file.filename):
        return jsonify({
            "status": "error",
            "message": "Unsupported file type. Please upload txt, pdf, doc, or docx."
        }), 400

    try:
        extracted_text = extract_text_from_file(file)

        if not extracted_text or not extracted_text.strip():
            return jsonify({
                "status": "error",
                "message": "Could not extract text from the uploaded file. Try a txt, text-based pdf, or docx file."
            }), 400

        prediction, result_text, confidence = \
    predict_legal(extracted_text)
        explanation = ""

        if use_xai:
            try:
                explanation = explain_legal_shap(extracted_text)
            except Exception as e:
                explanation = f"Explanation could not be generated: {str(e)}"

        email = request.form.get("email")
        if email:
            save_history(
                email,
                "Legal Analysis",
                extracted_text,
                result_text
            )

        return jsonify({
            "status": "success",
            "filename": file.filename,
            "extracted_text": extracted_text[:3000],
            "prediction": prediction,
            "result": result_text,
            "confidence": confidence,
            "explanation": explanation
        })

    except Exception as e:
        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500

@app.route('/history/<email>', methods=['GET'])
def get_history(email):

    try:

        cur = mysql.connection.cursor()

        cur.execute(
            """
            SELECT
            analysis_type,
            input_text,
            result,
            created_at

            FROM analysis_history

            WHERE user_email=%s

            ORDER BY created_at DESC
            """,
            (email,)
        )

        rows = cur.fetchall()

        cur.close()


        history=[]


        for row in rows:

            history.append({

                "type": row[0],

                "text": row[1],

                "result": row[2],

                "date": str(row[3])
            })


        return jsonify({

            "status":"success",

            "history":history
        })


    except Exception as e:

        return jsonify({

            "status":"error",

            "message":str(e)
        }),500
    

@csrf.exempt
@app.route('/send-otp', methods=['POST'])
@limiter.limit("5 per minute")
def send_otp():

    try:

        data = request.get_json()

        email = data.get('email', '').strip()
        purpose = data.get('purpose', 'registration')

        if not email:
            return jsonify({
                'status': 'error',
                'message': 'Email is required'
            }), 400

        otp = str(random.randint(100000, 999999))

        otp_storage[email] = {
            "otp": otp,
            "expires": datetime.now() + timedelta(minutes=5)
        }

        msg = Message(
            'VeriLex AI Verification Code',
            recipients=[email]
        )

        if purpose == "email_update":


            msg.html = f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>VeriLex AI - Email Update Verification</title>
</head>

<body style="margin:0;padding:0;background:#f4f6f9;font-family:Arial,sans-serif;">

<div style="max-width:650px;margin:auto;background:white;">

    <div style="
        background:linear-gradient(135deg,#1e3a8a,#2563eb);
        padding:40px;
        text-align:center;
        color:white;
    ">
        <h1 style="margin:0;font-size:34px;">
            VeriLex AI
        </h1>

        <p style="margin-top:12px;font-size:18px;">
            Email Change Verification
        </p>
    </div>

    <div style="padding:35px;color:#333;">

        <h2 style="margin-top:0;">
            Verify Your New Email Address
        </h2>

        <p>
            We received a request to update the email
            address associated with your VeriLex AI account.
        </p>

        <p>
            Please use the OTP below to confirm this change:
        </p>

        <div style="
            background:#eff6ff;
            border:2px dashed #2563eb;
            text-align:center;
            padding:22px;
            font-size:34px;
            font-weight:bold;
            letter-spacing:8px;
            color:#1e40af;
            margin:25px 0;
        ">
            {otp}
        </div>

        <p>
            This OTP is valid for
            <strong>5 minutes</strong>.
        </p>

        <p>
            If you did not request this email update,
            please ignore this email and your account
            will remain unchanged.
        </p>

    </div>

    <div style="
        background:#f1f5f9;
        text-align:center;
        padding:15px;
        color:#64748b;
        font-size:14px;
    ">
        © VeriLex AI
    </div>

</div>

</body>
</html>
"""

        else:

            msg.html = f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>VeriLex AI Verification</title>
</head>

<body style="margin:0;padding:0;background:#f4f6f9;font-family:Arial,sans-serif;">

<div style="max-width:650px;margin:auto;background:white;">

    <div style="
        background:linear-gradient(135deg,#1e3a8a,#2563eb);
        padding:40px;
        text-align:center;
        color:white;
    ">
        <h1 style="margin:0;font-size:34px;">
            VeriLex AI
        </h1>

        <p style="margin-top:12px;font-size:18px;">
            AI-Powered Legal Document Verification
        </p>
    </div>

    <div style="padding:35px;color:#333;">

        <h2 style="margin-top:0;">
            Welcome to VeriLex AI
        </h2>

        <p>
            Thank you for registering with VeriLex AI.
        </p>

        <p>
            Please use the following One-Time Password (OTP)
            to verify your email address:
        </p>

        <div style="
            background:#eff6ff;
            border:2px dashed #2563eb;
            text-align:center;
            padding:22px;
            font-size:34px;
            font-weight:bold;
            letter-spacing:8px;
            color:#1e40af;
            margin:25px 0;
        ">
            {otp}
        </div>

        <p>
            This OTP is valid for
            <strong>5 minutes</strong>.
        </p>

        <p>
            If you did not create an account,
            you can safely ignore this email.
        </p>

    </div>

    <div style="
        background:#f1f5f9;
        text-align:center;
        padding:15px;
        color:#64748b;
        font-size:14px;
    ">
        © VeriLex AI
    </div>

</div>

</body>
</html>
"""

        mail.send(msg)

        logging.info(f"OTP sent to {email}")

        return jsonify({
            'status': 'success',
            'message': 'OTP sent successfully'
        }), 200

    except Exception as e:

        logging.error(f"OTP ERROR: {str(e)}")

        return jsonify({
            'status': 'error',
            'message': 'Failed to send OTP'
        }), 500

@csrf.exempt
@app.route('/verify-otp', methods=['POST'])
@limiter.limit("10 per minute")
def verify_otp():

    data = request.get_json()

    email = data.get('email', '').strip()
    otp = data.get('otp', '').strip()

    stored_data = otp_storage.get(email)

    if not stored_data:
        return jsonify({
            'status': 'error',
            'message': 'OTP not found'
        }), 400

    if datetime.now() > stored_data["expires"]:

        del otp_storage[email]

        return jsonify({
            'status': 'error',
            'message': 'OTP expired'
        }), 400

    if stored_data["otp"] != otp:

        return jsonify({
            'status': 'error',
            'message': 'Invalid OTP'
        }), 400

    verified_emails.add(email)
    del otp_storage[email]
    logging.info(f"OTP verified: {email}")
    return jsonify({
        'status': 'success',
        'message': 'OTP verified successfully'
    }), 200
    
# ============================================
# 🔑 SEND RESET PASSWORD OTP
# ============================================

@csrf.exempt
@app.route('/send-reset-otp', methods=['POST'])
@limiter.limit("5 per minute")
def send_reset_otp():

    try:

        data = request.get_json()

        email = data.get('email', '').strip()

        if not email:

            return jsonify({
                "status": "error",
                "message": "Email is required"
            }), 400

        cur = mysql.connection.cursor()

        cur.execute(
            "SELECT * FROM users WHERE email=%s",
            (email,)
        )

        user = cur.fetchone()

        cur.close()

        if not user:

            return jsonify({
                "status": "error",
                "message": "No account found with this email"
            }), 404

        otp = str(random.randint(100000, 999999))

        reset_otp_storage[email] = {
            "otp": otp,
            "expires": datetime.now() + timedelta(minutes=5)
        }

        msg = Message(
            'VeriLex AI Password Reset OTP',
            recipients=[email]
        )

        msg.html = f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
body {{
    background-color:#f4f7fb;
    font-family:Arial,sans-serif;
    margin:0;
    padding:0;
}}

.container {{
    max-width:600px;
    margin:30px auto;
    background:#ffffff;
    border-radius:12px;
    overflow:hidden;
    box-shadow:0 4px 20px rgba(0,0,0,0.08);
}}

.header {{
    background:#2563eb;
    color:white;
    text-align:center;
    padding:25px;
}}

.header h1 {{
    margin:0;
}}

.content {{
    padding:30px;
    color:#333;
}}

.otp-box {{
    text-align:center;
    margin:30px 0;
}}

.otp {{
    display:inline-block;
    background:#eff6ff;
    color:#2563eb;
    font-size:34px;
    font-weight:bold;
    letter-spacing:8px;
    padding:15px 30px;
    border-radius:10px;
}}

.warning {{
    color:#dc2626;
    font-weight:bold;
}}

.footer {{
    text-align:center;
    padding:20px;
    color:#6b7280;
    font-size:13px;
}}
</style>
</head>

<body>

<div class="container">

<div class="header">
<h1>VeriLex AI</h1>
<p>Password Reset Verification</p>
</div>

<div class="content">

<p>Hello,</p>

<p>
We received a request to reset your VeriLex AI account password.
</p>

<div class="otp-box">
<div class="otp">{otp}</div>
</div>

<p>
This OTP will expire in <strong>5 minutes</strong>.
</p>

<p class="warning">
If you did not request a password reset, please ignore this email.
</p>

</div>

<div class="footer">
© 2026 VeriLex AI • Secure Legal Intelligence Platform
</div>

</div>

</body>
</html>
"""

        mail.send(msg)

        logging.info(f"Password reset OTP sent to {email}")

        return jsonify({
            "status": "success",
            "message": "Reset OTP sent successfully"
        }), 200

    except Exception as e:

        logging.error(f"RESET OTP ERROR: {str(e)}")

        return jsonify({
            "status": "error",
            "message": "Failed to send reset OTP"
        }), 500

# ============================================
# 🔒 RESET PASSWORD
# ============================================

@csrf.exempt
@app.route('/reset-password', methods=['POST'])
@limiter.limit("5 per minute")
def reset_password():

    try:

        data = request.get_json()

        email = data.get('email', '').strip()
        otp = data.get('otp', '').strip()
        new_password = data.get('newPassword', '').strip()

        if not email or not otp or not new_password:

            return jsonify({
                "status": "error",
                "message": "All fields are required"
            }), 400

        # PASSWORD VALIDATION
        if not is_valid_password(new_password):

            return jsonify({
                "status": "error",
                "message":
                "Password must contain minimum 10 characters, 2 uppercase letters, 2 lowercase letters, 2 numbers and 1 special character."
            }), 400

        stored_data = reset_otp_storage.get(email)

        if not stored_data:

            return jsonify({
                "status": "error",
                "message": "OTP not found"
            }), 400

        # OTP EXPIRY CHECK
        if datetime.now() > stored_data["expires"]:

            del reset_otp_storage[email]

            return jsonify({
                "status": "error",
                "message": "OTP expired"
            }), 400

        # OTP MATCH CHECK
        if stored_data["otp"] != otp:

            return jsonify({
                "status": "error",
                "message": "Invalid OTP"
            }), 400

        # HASH NEW PASSWORD
        hashed_password = bcrypt.hashpw(
            new_password.encode('utf-8'),
            bcrypt.gensalt()
        )

        cur = mysql.connection.cursor()

        cur.execute(
            """
            UPDATE users
            SET password=%s
            WHERE email=%s
            """,
            (
                hashed_password.decode('utf-8'),
                email
            )
        )

        mysql.connection.commit()

        cur.close()

        # REMOVE USED OTP
        del reset_otp_storage[email]

        logging.info(f"Password reset successful: {email}")

        return jsonify({
            "status": "success",
            "message": "Password reset successful. Please login again."
        }), 200

    except Exception as e:

        logging.error(f"RESET PASSWORD ERROR: {str(e)}")

        return jsonify({
            "status": "error",
            "message": "Password reset failed"
        }), 500

# ============================================
# 📩 CONTACT FORM API
# ============================================

@csrf.exempt
@app.route('/contact-message', methods=['POST'])
def contact_message():
    data = request.get_json()
    captcha_token = data.get("captcha")

    if not verify_recaptcha(captcha_token):
        return jsonify({
        "status": "error",
        "message": "CAPTCHA verification failed"
    }), 400
    name = data.get('name', '').strip()
    email = data.get('email', '').strip()
    subject = data.get('subject', '').strip()
    message = data.get('message', '').strip()

    if not name or not email or not subject or not message:
        return jsonify({
            "status": "error",
            "message": "All fields are required."
        }), 400

    try:
        cur = mysql.connection.cursor()
        cur.execute(
            """
            INSERT INTO contact_messages (name, email, subject, message)
            VALUES (%s, %s, %s, %s)
            """,
            (name, email, subject, message)
        )
        mysql.connection.commit()
        cur.close()

        return jsonify({
            "status": "success",
            "message": "Your message has been sent successfully."
        })

    except Exception as e:
        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500
    
@csrf.exempt
@app.route('/save-profile', methods=['POST'])
def save_profile():

    try:
        print("FORM DATA:", request.form)
        print("FILES:", request.files)

        email = request.form.get('email')

        if not email:
            return jsonify({
                "status": "error",
                "message": "Email is missing"
            }), 400

        username = request.form.get('username')
        full_name = request.form.get('full_name')
        gender = request.form.get('gender')
        dob = request.form.get('dob')
        print("DOB RECEIVED:", repr(dob))
        if not dob or dob.strip() == "":
            dob = None
        country = request.form.get('country')
        occupation = request.form.get('occupation')
        bio = request.form.get('bio')

        profile_picture = None

        # IMAGE UPLOAD
        if 'profile_picture' in request.files:

            file = request.files.get('profile_picture')

            if file and file.filename != '':

                if not allowed_image(file.filename):

                    return jsonify({
                        "status": "error",
                        "message": "Invalid image format"
                    }), 400

                filename = secure_filename(file.filename)

                # UNIQUE FILE NAME
                filename = (
                    str(datetime.now().timestamp())
                    + "_"
                    + filename
                )

                # filepath = os.path.join(
                #     app.config['PROFILE_UPLOAD_FOLDER'],
                #     filename
                # )

                # file.save(filepath)

                # # SAVE ONLY FILE NAME IN DATABASE
                # profile_picture = filename
                if bucket:

                    blob = bucket.blob(filename)

                    blob.upload_from_file(
                        file,
                        content_type=file.content_type
                    )

                else:

                    filepath = os.path.join(
                        app.config["PROFILE_UPLOAD_FOLDER"],
                        filename
                    )

                    file.save(filepath)

                profile_picture = filename

        cur = mysql.connection.cursor()

        # CHECK EXISTING PROFILE
        cur.execute(
            """
            SELECT * FROM user_profiles
            WHERE user_email=%s
            """,
            (email,)
        )

        existing = cur.fetchone()

        if existing:

            if profile_picture:

                cur.execute(
                    """
                    UPDATE user_profiles
                    SET
                    username=%s,
                    full_name=%s,
                    gender=%s,
                    dob=%s,
                    country=%s,
                    occupation=%s,
                    bio=%s,
                    profile_picture=%s

                    WHERE user_email=%s
                    """,
                    (
                        username,
                        full_name,
                        gender,
                        dob if dob else None,
                        country,
                        occupation,
                        bio,
                        profile_picture,
                        email
                    )
                )

            else:

                cur.execute(
                    """
                    UPDATE user_profiles
                    SET
                    username=%s,
                    full_name=%s,
                    gender=%s,
                    dob=%s,
                    country=%s,
                    occupation=%s,
                    bio=%s

                    WHERE user_email=%s
                    """,
                    (
                        username,
                        full_name,
                        gender,
                        dob if dob else None,
                        country,
                        occupation,
                        bio,
                        email
                    )
                )

        else:

            cur.execute(
                """
                INSERT INTO user_profiles(
                    user_email,
                    username,
                    full_name,
                    gender,
                    dob,
                    country,
                    occupation,
                    bio,
                    profile_picture
                )

                VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s)
                """,
                (
                    str(email),
                    str(username),
                    str(full_name),
                    str(gender),
                    dob if dob else None,
                    str(country),
                    str(occupation),
                    str(bio),
                    str(profile_picture) if profile_picture else None
                )
            )

        mysql.connection.commit()

        cur.close()

        return jsonify({
            "status": "success",
            "message": "Profile saved successfully"
        })

    except Exception as e:

        print("SAVE PROFILE ERROR:", str(e))

        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500
    
@app.route('/get-profile/<email>', methods=['GET'])
def get_profile(email):

    try:

        cur = mysql.connection.cursor()

        cur.execute(
            """
            SELECT *
            FROM user_profiles
            WHERE user_email=%s
            """,
            (email,)
        )

        row = cur.fetchone()

        if not row:

            cur.execute(
                """
                SELECT name,email
                FROM users
                WHERE email=%s
                """,
                (email,)
            )

            user = cur.fetchone()

            if user:

                cur.close()

                return jsonify({
                    "status": "success",
                    "profile": {
                        "email": user[1],
                        "username": user[0],
                        "full_name": user[0],
                        "gender": "",
                        "dob": "",
                        "country": "",
                        "occupation": "",
                        "bio": "",
                        "profile_picture": None
                    }
                })

            cur.close()

            return jsonify({
                "status": "success",
                "profile": None
            })

        profile = {

            "email": row[1],
            "username": row[2],
            "full_name": row[3],
            "gender": row[4],
            "dob": str(row[5]) if row[5] else "",
            "country": row[6],
            "occupation": row[7],
            "bio": row[8],
            "profile_picture":
                f"{os.getenv('BACKEND_URL')}/profile-picture/{row[9]}"
                if row[9] else None
        }

        cur.close()

        return jsonify({
            "status": "success",
            "profile": profile
        })

    except Exception as e:

        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500
    
@csrf.exempt
@app.route('/update-email', methods=['POST'])
def update_email():

    try:

        data = request.get_json()

        old_email = data.get('old_email')
        new_email = data.get('new_email')

        cur = mysql.connection.cursor()

        # Check old user exists
        cur.execute(
            "SELECT id, email FROM users WHERE email=%s",
            (old_email,)
        )

        user = cur.fetchone()

        print("FOUND USER =", user)

        if not user:
            return jsonify({
                "status": "error",
                "message": "Old email not found"
            }), 404

        # Check new email already exists
        cur.execute(
            "SELECT * FROM users WHERE email=%s",
            (new_email,)
        )

        existing = cur.fetchone()

        if existing:
            return jsonify({
                "status": "error",
                "message": "New email already exists"
            }), 400

        # Update users table
        cur.execute(
            """
            UPDATE users
            SET email=%s
            WHERE email=%s
            """,
            (new_email, old_email)
        )

        print("USERS UPDATED =", cur.rowcount)

        # Update profile table
        cur.execute(
            """
            UPDATE user_profiles
            SET user_email=%s
            WHERE user_email=%s
            """,
            (new_email, old_email)
        )

        print("PROFILE UPDATED =", cur.rowcount)

        mysql.connection.commit()

        # VERIFY
        cur.execute(
            "SELECT email FROM users WHERE email=%s",
            (new_email,)
        )

        verify = cur.fetchone()

        print("AFTER UPDATE =", verify)

        cur.close()

        return jsonify({
            "status": "success",
            "message": "Email updated successfully"
        })

    except Exception as e:

        print("EMAIL UPDATE ERROR =", e)

        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500
    
@csrf.exempt
@app.route('/change-password', methods=['POST'])
def change_password():

    try:

        data = request.get_json()

        email = data.get('email')
        old_password = data.get('old_password')
        new_password = data.get('new_password')

        cur = mysql.connection.cursor()

        cur.execute(
            """
            SELECT password
            FROM users
            WHERE email=%s
            """,
            (email,)
        )

        user = cur.fetchone()

        if not user:

            return jsonify({
                "status": "error",
                "message": "User not found"
            }), 404

        db_password = user[0]

        if not bcrypt.checkpw(
            old_password.encode('utf-8'),
            db_password.encode('utf-8')
        ):

            return jsonify({
                "status": "error",
                "message": "Old password incorrect"
            }), 401

        if not is_valid_password(new_password):

            return jsonify({
        "status": "error",
        "message": "Password does not meet security requirements"
    }), 400
        
        hashed_password = bcrypt.hashpw(
            new_password.encode('utf-8'),
            bcrypt.gensalt()
        )

        cur.execute(
            """
            UPDATE users
            SET password=%s
            WHERE email=%s
            """,
            (
                hashed_password.decode('utf-8'),
                email
            )
        )

        mysql.connection.commit()

        cur.close()

        return jsonify({
            "status": "success",
            "message": "Password updated successfully"
        })

    except Exception as e:

        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500

from flask import send_file

@app.route("/profile-picture/<filename>")
def profile_picture(filename):

    try:

        blob = bucket.blob(filename)

        if not blob.exists():

            return jsonify({
                "status":"error",
                "message":"Image not found"
            }),404

        image_bytes = blob.download_as_bytes()

        return send_file(
            io.BytesIO(image_bytes),
            mimetype=blob.content_type
        )

    except Exception as e:

        return jsonify({
            "status":"error",
            "message":str(e)
        }),500

@csrf.exempt
@app.route('/delete-profile-picture', methods=['POST'])
def delete_profile_picture():

    try:

        data = request.get_json()

        email = data.get('email')

        if not email:
            return jsonify({
                "status": "error",
                "message": "Email is required"
            }), 400

        cur = mysql.connection.cursor()

        # Get existing image filename
        cur.execute(
            """
            SELECT profile_picture
            FROM user_profiles
            WHERE user_email=%s
            """,
            (email,)
        )

        row = cur.fetchone()

        if not row:
            cur.close()

            return jsonify({
        "status": "success",
        "message": "No profile picture found"
    })

        filename = row[0]

        if not filename:

            cur.execute("""
            UPDATE user_profiles
            SET profile_picture=NULL
            WHERE user_email=%s
        """, (email,))

            mysql.connection.commit()
            cur.close()

            return jsonify({
            "status": "success",
            "message": "Profile picture deleted successfully"
    })

        # Delete image from Google Cloud Storage

        if filename:

            if bucket:

                blob = bucket.blob(filename)

                if blob.exists():

                    blob.delete()

            else:

                filepath = os.path.join(
                    app.config["PROFILE_UPLOAD_FOLDER"],
                    filename
                )

                if os.path.exists(filepath):

                    os.remove(filepath)

        # Remove filename from database
        cur.execute(
            """
            UPDATE user_profiles
            SET profile_picture=NULL
            WHERE user_email=%s
            """,
            (email,)
        )

        mysql.connection.commit()

        cur.close()

        return jsonify({
            "status": "success",
            "message": "Profile picture deleted successfully"
        })

    except Exception as e:

        print("DELETE PROFILE PIC ERROR:", str(e))

        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500

@csrf.exempt
@app.route('/track-visitor', methods=['POST'])
def track_visitor():

    data = request.get_json() or {}

    device_id = data.get("deviceId")
    email = data.get("email")

    ip_address = request.remote_addr

    if not device_id:
        return jsonify({
            "status": "error",
            "message": "Missing device id"
        }), 400

    if email:

        visitor_key = f"{device_id}_{email}"

    else:

        visitor_key = device_id

    try:

        cur = mysql.connection.cursor()

        cur.execute(
            """
            SELECT id
            FROM visitors
            WHERE visitor_key=%s
            """,
            (visitor_key,)
        )

        existing = cur.fetchone()

        if not existing:

            cur.execute(
                """
                INSERT INTO visitors
                (
                    visitor_key,
                    email,
                    device_id,
                    ip_address
                )
                VALUES(%s,%s,%s,%s)
                """,
                (
                    visitor_key,
                    email,
                    device_id,
                    ip_address
                )
            )

            mysql.connection.commit()

        cur.execute(
            "SELECT COUNT(*) FROM visitors"
        )

        total = cur.fetchone()[0]

        cur.close()

        return jsonify({
            "status": "success",
            "totalVisitors": total
        })

    except Exception as e:

        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500
    
@csrf.exempt
@app.route('/visitor-count', methods=['GET'])
def visitor_count():

    try:

        cur = mysql.connection.cursor()

        cur.execute(
            "SELECT COUNT(*) FROM visitors"
        )

        total = cur.fetchone()[0]

        cur.close()

        return jsonify({
            "totalVisitors": total
        })

    except Exception as e:

        return jsonify({
            "message": str(e)
        }), 500
    
# ============================================
# ⚖ CONSTITUTIONAL Q&A API
# ============================================

@csrf.exempt
@app.route('/ask-constitution', methods=['POST'])
def ask_constitution():

    try:

        data = request.get_json()

        question = data.get(
            "question",
            ""
        ).strip()

        email = data.get(
            "email"
        )

        key = os.getenv(
            "GEMINI_API_KEY"
        )

        genai.configure(
            api_key=key
        )

        model = genai.GenerativeModel(
            "gemini-2.5-flash"
        )

        response = model.generate_content(
            f"""
Answer the following constitutional question clearly and professionally:

{question}
"""
        )

        answer = response.text

        save_history(
            email,
            "Constitutional Q&A",
            question,
            answer[:500]
        )

        return jsonify({
            "answer": answer
        })

    except Exception as e:

        print(
            "CONSTITUTION ERROR:",
            str(e)
        )

        return jsonify({
            "error": str(e)
        }), 500
    
@csrf.exempt
@app.route('/constitution-upload', methods=['POST'])
def constitution_upload():

    try:

        if 'file' not in request.files:

            return jsonify({
                "message": "No file uploaded."
            }), 400

        file = request.files['file']

        if file.filename == '':

            return jsonify({
                "message": "No file selected."
            }), 400

        email = request.form.get("email")

        filename = file.filename.lower()

        extracted_text = ""

        # TXT
        if filename.endswith('.txt'):

            extracted_text = file.read().decode(
                'utf-8',
                errors='ignore'
            )

        # PDF
        elif filename.endswith('.pdf'):

            pdf = fitz.open(
                stream=file.read(),
                filetype="pdf"
            )

            for page in pdf:

                extracted_text += (
                    page.get_text() + "\n"
                )

        # DOCX
        elif filename.endswith('.docx'):

            doc = docx.Document(file)

            extracted_text = "\n".join(
                p.text
                for p in doc.paragraphs
            )

        else:

            return jsonify({
                "message":
                "Only TXT, PDF and DOCX files are supported."
            }), 400

        # Prevent huge prompts
        extracted_text = extracted_text[:20000]

        # Gemini
        key = os.getenv(
            "GEMINI_API_KEY"
        )

        genai.configure(
            api_key=key
        )

        model = genai.GenerativeModel(
            "gemini-2.5-flash"
        )

        response = model.generate_content(
            f"""
Answer the following constitutional question clearly and professionally:

{extracted_text}
"""
        )

        answer = response.text

        save_history(
            email,
            "Constitutional Q&A",
            extracted_text[:5000],
            answer[:500]
        )

        return jsonify({

            "extracted_text":
            extracted_text,

            "answer":
            answer

        })

    except Exception as e:

        print(
            "CONSTITUTION UPLOAD ERROR:",
            str(e)
        )

        return jsonify({
            "message": str(e)
        }), 500
    
from routes.constitution import constitution_bp

app.register_blueprint(
    constitution_bp
)

@csrf.exempt
@app.route("/google-login", methods=["POST"])
def google_login():

    try:

        data = request.get_json()

        name = data.get("name")
        email = data.get("email")
        google_id = data.get("google_id", "")

        if not email:

            return jsonify({
                "message": "Email is required"
            }), 400

        cursor = mysql.connection.cursor()

        cursor.execute(
            """
            SELECT id, name, email, is_active
            FROM users
            WHERE email=%s
            """,
            (email,)
        )

        user = cursor.fetchone()

        if user:

            if user[3] == 0:

                return jsonify({
                    "message":
                    "Your account has been disabled by the administrator."
                }), 403
            
            cursor.execute(
                """
                UPDATE users
                SET last_login = NOW()
                WHERE email=%s
                """,
                (email,)
            )

            mysql.connection.commit()

            cursor.execute("""
            SELECT
            username,
            gender,
            dob,
            country,
            occupation,
            bio
            FROM user_profiles
            WHERE user_email=%s
            """, (email,))

            profile = cursor.fetchone()

            if not profile:

                profile_completed = False

            else:

                username, gender, dob, country, occupation, bio = profile

                profile_completed = all([
                    username,
                    gender,
                    dob,
                    country,
                    occupation,
                    bio
                ])
            cursor.close()
            return jsonify({
                "message": "Login successful",
                "name": user[1],
                "email": user[2],
                "profile_completed": profile_completed
            }), 200

        else:

            cursor.execute(
                """
                INSERT INTO users
                (
                    name,
                    email,
                    password,
                    google_id,
                    login_provider
                )
                VALUES
                (
                    %s,
                    %s,
                    %s,
                    %s,
                    %s
                )
                """,
                (
                    name,
                    email,
                    None,
                    google_id,
                    "google"
                )
            )

            mysql.connection.commit()

            return jsonify({
                "message": "Google account created",
                "name": name,
                "email": email,
                "profile_completed": False
            }), 201

    except Exception as e:

        print(
            "Google Login Error:",
            str(e)
        )

        return jsonify({
            "message": str(e)
        }), 500

    finally:

        try:
            cursor.close()
        except:
            pass

@csrf.exempt
@app.route('/admin/login', methods=['POST'])
def admin_login():

    data = request.get_json()

    username = data.get("username")
    password = data.get("password")

    cur = mysql.connection.cursor()

    cur.execute(
        """
        SELECT password
        FROM admin_users
        WHERE username=%s
        """,
        (username,)
    )

    admin = cur.fetchone()

    cur.close()

    if not admin:

        return jsonify({
            "status":"error",
            "message":"Invalid credentials"
        }),401

    if password != admin[0]:

        return jsonify({
            "status":"error",
            "message":"Invalid credentials"
        }),401

    return jsonify({
        "status":"success"
    })

@csrf.exempt
@app.route('/admin/send-notification', methods=['POST'])
def send_notification():

    data = request.get_json()

    message = data.get("message")
    if not message:

        return jsonify({
            "status":"error",
            "message":"Notification message cannot be empty."
        }), 400

    cur = mysql.connection.cursor()

    cur.execute(
        """
        INSERT INTO notifications(message)
        VALUES(%s)
        """,
        (message,)
    )

    mysql.connection.commit()
    add_log(
    "Notification Sent",
    message
)
    
    cur.close()

    return jsonify({
        "status":"success",
        "message":"Notification sent"
    })

@app.route('/notifications', methods=['GET'])
def get_notifications():

    cur = mysql.connection.cursor()

    cur.execute(
        """
        SELECT id,message,created_at
        FROM notifications
        ORDER BY id DESC
        """
    )

    rows = cur.fetchall()

    cur.close()

    data = []

    for row in rows:

        data.append({
            "id":row[0],
            "message":row[1],
            "time":str(row[2])
        })

    return jsonify(data)

@csrf.exempt
@app.route('/admin/delete-notification/<int:id>',
methods=['DELETE'])
def delete_notification(id):

    cur = mysql.connection.cursor()

    cur.execute(
        """
        DELETE FROM notifications
        WHERE id=%s
        """,
        (id,)
    )

    mysql.connection.commit()

    cur.close()

    return jsonify({
        "status":"success"
    })

@csrf.exempt
@app.route('/admin/clear-notifications',
methods=['DELETE'])
def clear_notifications():

    cur = mysql.connection.cursor()

    cur.execute(
        """
        DELETE FROM notifications
        """
    )

    mysql.connection.commit()

    cur.close()

    return jsonify({
        "status":"success"
    })

@app.route('/admin/user-count')
def user_count():

    cur = mysql.connection.cursor()

    cur.execute(
        "SELECT COUNT(*) FROM users"
    )

    total = cur.fetchone()[0]

    cur.close()

    return jsonify({
        "total_users":total
    })

@app.route('/admin/users')
def admin_users():

    cur = mysql.connection.cursor()

    cur.execute(
        """
        SELECT id,name,email, is_active, created_at, last_login
        FROM users
        """
    )

    users = cur.fetchall()

    cur.close()

    return jsonify(users)

@csrf.exempt
@app.route('/admin/delete-user/<id>',
methods=['DELETE'])
def delete_user(id):

    cur = mysql.connection.cursor()

    cur.execute(
        "DELETE FROM users WHERE id=%s",
        (id,)
    )

    mysql.connection.commit()
    add_log(
    "Delete User",
    f"User ID {id} deleted"
)
    
    cur.close()

    return jsonify({
        "status":"success"
    })

@app.route('/admin/contact-messages')
def admin_messages():

    cur = mysql.connection.cursor()

    cur.execute(
        """
        SELECT name,email,subject,message
        FROM contact_messages
        ORDER BY id DESC
        """
    )

    data = cur.fetchall()

    cur.close()

    return jsonify(data)

@csrf.exempt
@app.route('/admin/notification-count')
def notification_count():

    cur = mysql.connection.cursor()

    cur.execute(
        """
        SELECT COUNT(*)
        FROM notifications
        """
    )

    total = cur.fetchone()[0]

    cur.close()

    return jsonify({
        "total_notifications": total
    })

@csrf.exempt
@app.route('/admin/contact-count')
def contact_count():

    cur = mysql.connection.cursor()

    cur.execute(
        """
        SELECT COUNT(*)
        FROM contact_messages
        """
    )

    total = cur.fetchone()[0]

    cur.close()

    return jsonify({
        "total_contacts": total
    })

@csrf.exempt
@app.route('/admin/maintenance', methods=['GET'])
def get_maintenance():

    try:

        cur = mysql.connection.cursor()

        cur.execute("""
            SELECT maintenance
            FROM settings
            WHERE id=1
        """)

        result = cur.fetchone()

        cur.close()

        if result is None:

            return jsonify({
                "maintenance": False
            })

        if isinstance(result, dict):

            return jsonify({
                "maintenance":
                result.get("maintenance", False)
            })

        return jsonify({
            "maintenance": result[0]
        })

    except Exception as e:

        print("MAINTENANCE ERROR:", str(e))

        return jsonify({
            "maintenance": False
        })

@csrf.exempt
@app.route('/admin/maintenance', methods=['POST'])
def set_maintenance():

    data = request.get_json()

    value = data.get("maintenance")

    cur = mysql.connection.cursor()

    cur.execute(
        """
        UPDATE settings
        SET maintenance=%s
        WHERE id=1
        """,
        (value,)
    )

    mysql.connection.commit()
    if value:

        add_log(
        "Maintenance Enabled",
        "Website put into maintenance mode"
    )

    else:

        add_log(
        "Maintenance Disabled",
        "Website restored for public access"
    )
    cur.close()

    return jsonify({
        "status":"success"
    })

@app.route('/admin/all-notifications', methods=['GET'])
def admin_all_notifications():

    cur = mysql.connection.cursor()

    cur.execute(
        """
        SELECT
        id,
        message,
        created_at
        FROM notifications
        ORDER BY id DESC
        """
    )

    data = cur.fetchall()

    cur.close()

    return jsonify(data)

@app.route('/admin/delete-notification/<int:id>', methods=['DELETE'])
def admin_delete_notification(id):

    cur = mysql.connection.cursor()

    cur.execute(
        """
        DELETE FROM notifications
        WHERE id=%s
        """,
        (id,)
    )

    mysql.connection.commit()

    cur.close()

    return jsonify({
        "status":"success"
    })

@csrf.exempt
@app.route('/admin/toggle-user-status/<int:user_id>', methods=['POST'])
def toggle_user_status(user_id):

    cur = mysql.connection.cursor()

    cur.execute(
        """
        SELECT is_active
        FROM users
        WHERE id=%s
        """,
        (user_id,)
    )

    user = cur.fetchone()

    if not user:

        cur.close()

        return jsonify({
            "status":"error"
        }),404

    new_status = not bool(user[0])

    cur.execute(
        """
        UPDATE users
        SET is_active=%s
        WHERE id=%s
        """,
        (
            new_status,
            user_id
        )
    )

    mysql.connection.commit()
    add_log(
    "User Status Changed",
    f"Status changed for User ID {user_id}"
)
    cur.close()

    return jsonify({
        "status":"success"
    })

@app.route("/admin/activity-logs")
def activity_logs():

    cursor = mysql.connection.cursor()

    cursor.execute(
        """
        SELECT
        id,
        action,
        description,
        created_at
        FROM activity_logs
        ORDER BY id DESC
        """
    )

    logs = cursor.fetchall()

    cursor.close()

    return jsonify(logs)

@app.route("/admin/user-profile/<email>")
def get_user_profile(email):

    cur = mysql.connection.cursor()

    cur.execute(
        """
        SELECT
        username,
        full_name,
        gender,
        dob,
        country,
        occupation,
        bio,
        profile_picture,
        created_at,
        updated_at
        FROM user_profiles
        WHERE user_email=%s
        """,
        (email,)
    )

    profile = cur.fetchone()
    
    cur.close()

    if not profile:

        return jsonify({
            "message":"Profile not found"
        }),404

    return jsonify({
        "username": profile[0],
        "full_name": profile[1],
        "gender": profile[2],
        "dob": str(profile[3]) if profile[3] else "",
        "country": profile[4],
        "occupation": profile[5],
        "bio": profile[6],
        "profile_picture": profile[7],
        "created_at": str(profile[8]),
        "updated_at": str(profile[9])
    })

# ============================================
# ▶ RUN SERVER
# ============================================

if __name__ == '__main__':
    app.run(debug=False)